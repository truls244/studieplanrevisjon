from flask import Blueprint, jsonify, send_file, make_response
from sqlalchemy.orm import joinedload
from app.models import Studyprogram, Studyplan, Course, Institute, Semester, SemesterCourses
from io import BytesIO
from docx import Document
import datetime
import math
from docx.shared import Inches, Pt
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls
from docx.enum.text import WD_ALIGN_PARAGRAPH
from services import ServiceFactory

# Create a Blueprint for export routes
exportdocx_bp = Blueprint('exportdocx', __name__)

def calculate_year(studyplan_year, semester_number, term):
    return studyplan_year + (semester_number - 1) // 2 + (1 if term == "V" else 0)

def generate_studyplan_docx(studyprogram, studyplans):
    doc = Document()

    doc.add_heading(f"{studyprogram.name}", level=1)
    doc.add_paragraph(f"Institution: {studyprogram.institute.name}")
    doc.add_paragraph(f"Ansvarlig: {studyprogram.institute.ansvarlig}")
    doc.add_paragraph(f"Degree Type: {studyprogram.degree_type}")
    doc.add_paragraph("\n")

    # start-slutt år
    first_studyplan = studyplans[0]
    start_year = first_studyplan['year']
    end_year = calculate_year(start_year, first_studyplan['semesters'][-1]['semester_number'], first_studyplan['semesters'][-1]['term'])
    # doc.add_paragraph(f"Period: {start_year} - {end_year}")
    doc.add_heading(f"Studieplanmatrise", level=2)
    doc.add_paragraph(f"Studieplan {start_year} - {end_year} for {studyprogram.name}", style="Normal")

    table = doc.add_table(rows=len(first_studyplan['semesters']) + 1, cols=4)
    table.style = 'Table Grid'

    header_cells = table.rows[0].cells
    header_cells[0].text = "Semester"
    header_cells[1].text = "10 sp"
    header_cells[2].text = "10 sp"
    header_cells[3].text = "10 sp"

    # NUMMER 1
    for row_idx, semester in enumerate(first_studyplan['semesters'], start=1):
        row = table.rows[row_idx].cells
        semester_year = calculate_year(start_year, semester['semester_number'], semester['term'])

        row[0].text = f"{semester['semester_number']} ({semester['term']}-{semester_year})"

        

        blocks = ["", "", ""]  
        current_block = 0
        current_block_credits = 0

        for course in semester['semester_courses']:
            if course['is_elective']:
                blocks = ["VALGEMNE", "VALGEMNE", "VALGEMNE"]
                break  
            else:
                course_text = f"{course['courseCode']} ({course['credits']} sp)"
                course_credits = course['credits'] if course['credits'] > 0 else 0


                if current_block_credits + course_credits <= 10:
                    blocks[current_block] += course_text + "\n"
                    current_block_credits += course_credits
                else:

                    current_block += 1
                    if current_block < 3:  
                        blocks[current_block] += course_text + "\n"
                        current_block_credits = course_credits


        for col_idx in range(3):
            row[col_idx + 1].text = blocks[col_idx] if blocks[col_idx] else ""

    doc.add_paragraph("\n")

    # NUMMER 1
    doc.add_heading("Emneoversikt", level=2)
    doc.add_paragraph(f"Studieplan {start_year} - {end_year} for {studyprogram.name}", style="Normal")

    for semester in first_studyplan['semesters']:
        semester_year = calculate_year(start_year, semester['semester_number'], semester['term'])
        doc.add_heading(f"Semester {semester['semester_number']} ({semester['term']} {semester_year})", level=3)


        electives_by_category = {}
        for course in semester['semester_courses']:
            if course['is_elective']:
                category_name = course['category']['name'] if course.get('category') else "Uncategorized"
                if category_name not in electives_by_category:
                    electives_by_category[category_name] = []
                electives_by_category[category_name].append(course)
            else:

                doc.add_paragraph(
                    f"{course['courseCode']} - {course['name']} ({course['credits']} sp)",
                    style="Normal"
                )

        for category, courses in electives_by_category.items():
            doc.add_heading(f"{category}", level=4)
            for course in courses:
                doc.add_paragraph(
                    f"{course['courseCode']} - {course['name']} ({course['credits']} sp)",
                    style="Normal"
                )

    doc.add_paragraph("\n")
    # EKSTRA studieplana. 
    for studyplan_idx, studyplan in enumerate(studyplans[1:], start=2):
        doc.add_paragraph("\n")

        if studyplan_idx == 2:
            semesters_to_skip = 2 # start from år 2?
        elif studyplan_idx == 3:
            semesters_to_skip= 4  # start fra år 4?
        else:
            continue 

        organized_semesters = studyplan['semesters'][semesters_to_skip:]
        start_year = calculate_year(studyplan['year'], semesters_to_skip + 1, "H")
        end_year = calculate_year(studyplan['year'], organized_semesters[-1]['semester_number'], organized_semesters[-1]['term'])
        doc.add_heading("Emneoversikt", level=2)
        doc.add_paragraph(f"Studieplan {start_year} - {end_year} for {studyprogram.name}", style="Normal")

        # fag per kull

        for semester in organized_semesters:
            semester_year = calculate_year(studyplan['year'], semester['semester_number'], semester['term'])
            doc.add_heading(f"Semester {semester['semester_number']} ({semester['term']}, {semester_year})", level=3)

            electives_by_category = {}
            for course in semester['semester_courses']:
                if course['is_elective']:
                    category_name = course['category']['name'] if course.get('category') else "Uncategorized"
                    if category_name not in electives_by_category:
                        electives_by_category[category_name] = []
                    electives_by_category[category_name].append(course)
                else:
                    doc.add_paragraph(
                        f"{course['courseCode']} - {course['name']} ({course['credits']} sp)",
                        style="Normal"
                    )

            for category, courses in electives_by_category.items():
                doc.add_heading(f"{category}", level=4)
                for course in courses:
                    doc.add_paragraph(
                        f"{course['courseCode']} - {course['name']} ({course['credits']} sp)",
                        style="Normal"
                    )

    return doc



@exportdocx_bp.route('/<int:studyprogram_id>', methods=['GET'])
def export_to_docx(studyprogram_id):
    try:
        studyplan_service = ServiceFactory.get_studyplan_service()
        studyprogram_service = ServiceFactory.get_studyprogram_service()


        studyprogram = studyprogram_service.get_studyprogram_by_id(studyprogram_id)
        if not studyprogram:
            return jsonify({"error": f"Study program with ID {studyprogram_id} not found"}), 404


        studyplans = studyplan_service.get_plans_for_export(studyprogram_id)

        doc = generate_studyplan_docx(studyprogram, studyplans)
        print("Document generated successfully", doc)


        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)


        filename = f"studyplan_{studyprogram_id}_{studyplans[0]['year']}.docx"


        return send_file(
            buffer,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

    except Exception as e:
        print(f"Error exporting study plan: {str(e)}")
        return jsonify({"error": str(e)}), 500